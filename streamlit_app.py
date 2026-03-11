import streamlit as st
import pandas as pd
from openai import OpenAI
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import os

# --- 1. PAGE SETUP ---
st.set_page_config(page_title="The Empowered PM Auditor", layout="wide", page_icon="🛡️")

# CSS: Styling for Action Buttons, Branding, and Sidebar
st.markdown("""
    <style>
    .stButton > button { width: 100%; border-radius: 8px; height: 3.5em; font-weight: bold; color: white !important; }
    
    /* Action Row Colors */
    [data-testid="stHorizontalBlock"] div:nth-child(1) button { background-color: #d32f2f !important; border: none; }
    [data-testid="stHorizontalBlock"] div:nth-child(2) button { background-color: #2e7d32 !important; border: none; }
    [data-testid="stHorizontalBlock"] div:nth-child(3) button { 
        background-color: #bdc3c7 !important; color: #000000 !important; border: 2px solid #2c3e50 !important;
    }

    .hero-section { background-color: #f0f4f8; padding: 25px; border-radius: 15px; border-left: 10px solid #0052cc; margin-bottom: 20px; }
    .cta-box { background-color: #fff3cd; padding: 20px; border-radius: 10px; border: 2px solid #ffeeba; margin-top: 25px; text-align: center; }
    .sidebar-guide { font-size: 0.85rem; color: #333; background-color: #f0f2f6; padding: 12px; border-radius: 8px; border: 1px solid #d1d5db; }
    .sample-table { font-size: 0.75rem; width: 100%; border-collapse: collapse; margin-top: 10px; background-color: white; }
    .sample-table th, .sample-table td { border: 1px solid #ccc; padding: 4px; text-align: left; color: black; }
    </style>
    """, unsafe_allow_html=True)

# --- 2. SECRETS & LOGIC ---
api_key = st.secrets.get("OPENAI_API_KEY", "")

if 'uploader_key' not in st.session_state:
    st.session_state['uploader_key'] = 0

def format_clean_text(doc, text):
    lines = text.split('\n')
    for line in lines:
        clean_line = line.replace('*', '').strip()
        if not clean_line:
            doc.add_paragraph()
            continue
        if line.strip().startswith('*') or line.strip().startswith('-'):
            doc.add_paragraph(clean_line, style='List Bullet')
        else:
            doc.add_paragraph(clean_line)

def create_word_doc(domain, audit, recovery):
    doc = Document()
    if os.path.exists("empPMlogo.png"):
        doc.add_picture("empPMlogo.png", width=Inches(1.2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    title = doc.add_heading('The Empowered PM: Audit & Recovery Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    tagline = doc.add_paragraph('Effective PM Practices for Everyone')
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tagline.runs[0].italic = True
    
    doc.add_heading('🕵️ Audit Findings', level=1)
    format_clean_text(doc, audit)
    
    doc.add_heading('🛠️ Recovery Roadmap', level=1)
    format_clean_text(doc, recovery)
    
    section = doc.sections[0]
    footer = section.footer
    f_p = footer.paragraphs[0]
    f_p.text = "Prepared by The Empowered PM Consulting, copyright 2026."
    f_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = f_p.runs[0] if f_p.runs else f_p.add_run()
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- 3. MAIN INTERFACE HEADER ---
col_logo, col_title = st.columns([1, 4])
with col_logo:
    if os.path.exists("empPMlogo.png"):
        st.image("empPMlogo.png", width=150)
with col_title:
    st.title("The Empowered PM Auditor")
    st.write("### *Effective PM Practices for Everyone*")

st.markdown("""
<div class="hero-section">
    <strong>Validation & Practical Recovery.</strong> Identify schedule errors and dependency gaps 
    to lead your project with absolute clarity and authority.
</div>
""", unsafe_allow_html=True)

# --- 4. SIDEBAR ---
with st.sidebar:
    st.header("Auditor Control")
    project_context = st.selectbox("Project Domain", ["IT/Software", "Construction", "Operations", "Marketing", "Events"])
    
    df_template = pd.DataFrame({
        "Task ID": [1, 2, 3, 4],
        "Task Name": ["Kickoff", "Design", "Execution", "Review"],
        "Start Date": ["2026-05-01", "2026-05-02", "2026-05-15", "2026-06-01"],
        "End Date": ["2026-05-01", "2026-05-10", "2026-05-30", "2026-06-02"],
        "Resource": ["PM", "Analyst", "Lead", "Client"],
        "Predecessor": ["None", "1", "2", "3"]
    })
    st.download_button(
        label="📥 Download CSV Template",
        data=df_template.to_csv(index=False).encode('utf-8'),
        file_name="empowered_pm_template.csv",
        mime="text/csv"
    )

    st.divider()
    st.markdown("**📖 Quick Start Guide**")
    st.markdown(f"""
    <div class="sidebar-guide">
    1. <b>File Prep:</b> Use our template for core columns.<br>
    2. <b>Select Type:</b> Set your <b>Project Domain</b>.<br><br>
    3. <b>Process:</b> Run Audit (Red), then Recovery (Green).<br><br>
    4. <b>Export:</b> Download your report.<br><br>
    5. <b>Reset:</b> Clear data for your next project.
    </div>
    """, unsafe_allow_html=True)

# --- 5. MAIN INTERFACE LOGIC ---
if not api_key:
    st.error("🚨 API Key missing from Secrets.")
    st.stop()

uploaded_file = st.file_uploader(
    "Upload Project Schedule (XLSX or CSV)", 
    type=["xlsx", "csv"], 
    key=f"uploader_{st.session_state['uploader_key']}"
)

if uploaded_file:
    client = OpenAI(api_key=api_key)
    try:
        df = pd.read_excel(uploaded_file) if uploaded_file.name.endswith('.xlsx') else pd.read_csv(uploaded_file)
        st.subheader("📋 Active Schedule Data")
        st.dataframe(df, use_container_width=True)
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if st.button("🚀 EXECUTE LOGIC AUDIT"):
                with st.spinner("Auditing logic..."):
                    schedule_data = df.to_string(index=False)
                    response = client.chat.completions.create(
                        model="gpt-4-turbo",
                        messages=[{"role": "system", "content": "You are a Senior PMO Auditor. Audit for logic errors. No tool suggestions. Use plain text math only."}, {"role": "user", "content": schedule_data}]
                    )
                    st.session_state['audit_report'] = response.choices[0].message.content
                    st.rerun()
        
        with col2:
            if 'audit_report' in st.session_state:
                if st.button("🛠️ GENERATE RECOVERY PLAN"):
                    with st.spinner("Building roadmap..."):
                        response = client.chat.completions.create(
                            model="gpt-4-turbo",
                            messages=[{"role": "system", "content": "Create a 3-step recovery roadmap based on findings. No asterisks."}, {"role": "user", "content": st.session_state['audit_report']}]
                        )
                        st.session_state['recovery_plan'] = response.choices[0].message.content
                        st.rerun()
        
        with col3:
            if st.button("🗑️ RESET ALL DATA"):
                if 'audit_report' in st.session_state: del st.session_state['audit_report']
                if 'recovery_plan' in st.session_state: del st.session_state['recovery_plan']
                st.session_state['uploader_key'] += 1
                st.rerun()

        if 'audit_report' in st.session_state:
            st.divider()
            st.subheader("🕵️ Auditor's Findings")
            st.markdown(st.session_state['audit_report'])
            
        if 'recovery_plan' in st.session_state:
            st.success("### ✅ Your Empowered Recovery Roadmap")
            st.markdown(st.session_state['recovery_plan'])
            
            # --- CUSTOMIZABLE CTA ---
            # Replace 'REPLACE_WITH_YOUR_SCHEDULE_LINK' with your Calendly/Booking link
            # Replace 'REPLACE_WITH_YOUR_LANDING_PAGE' with your website/program link
            st.markdown("""
                <div class="cta-box">
                    <h3 style="margin-top: 0; color: #856404;">🛠️ Need a Schedule Rescue?</h3>
                    <p style="color: #856404;">Don't present a broken plan. Book a <b>15-Minute Strategy Session</b> or join the <b>4-Week PM Survival Intensive</b>.</p>
                    <div style="display: flex; gap: 10px; justify-content: center;">
                        <a href="REPLACE_WITH_YOUR_SCHEDULE_LINK" target="_blank" style="text-decoration: none; flex: 1;">
                            <div style="background-color: #856404; color: white; padding: 12px; border-radius: 5px; font-weight: bold;">
                                Book a Rescue Session
                            </div>
                        </a>
                        <a href="REPLACE_WITH_YOUR_LANDING_PAGE" target="_blank" style="text-decoration: none; flex: 1;">
                            <div style="background-color: #2c3e50; color: white; padding: 12px; border-radius: 5px; font-weight: bold;">
                                Join the 4-Week Intensive
                            </div>
                        </a>
                    </div>
                </div>
            """, unsafe_allow_html=True)
            
            word_data = create_word_doc(project_context, st.session_state['audit_report'], st.session_state['recovery_plan'])
            st.download_button(label="📥 Download Report", data=word_data, file_name="Empowered_PM_Report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
            
    except Exception as e:
        st.error(f"Error: {e}")

st.divider()
with st.expander("📚 The Empowered PM's Glossary"):
    st.markdown("""<div class="glossary-card"><strong>Logic Validation:</strong> Ensuring your plan follows scheduling principles so it remains predictable and manageable.</div>""", unsafe_allow_html=True)
